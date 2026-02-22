from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import datetime
import io
import base64
import os
import tempfile

class DailyReportWordExporter:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """Настройка стилей документа"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Стиль заголовков
        heading_style = self.doc.styles['Heading 1']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
    
    def add_header(self, report_data):
        """Добавление шапки отчета"""
        title = self.doc.add_heading('ЕЖЕДНЕВНЫЙ ОТЧЕТ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = self.doc.add_paragraph('Технический надзор заказчика')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)
        
        self.doc.add_paragraph()
        
        # Основная информация в виде таблицы
        info_table = self.doc.add_table(rows=5, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('Дата отчета:', report_data.get('report_date', '')),
            ('Проект:', report_data.get('project_name', '')),
            ('Местоположение:', report_data.get('location', '')),
            ('Автор отчета:', report_data.get('author', '')),
            ('Погодные условия:', report_data.get('weather', ''))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = str(value)
            info_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    
    def add_equipment_section(self, form_data):
        """Добавление секции с техникой"""
        self.doc.add_heading('Задействованность техники на участке работ', level=1)
        
        # Получаем данные о технике из массивов
        equipment_names = form_data.getlist('equipment_name[]') if hasattr(form_data, 'getlist') else []
        damba_data = form_data.getlist('damba[]') if hasattr(form_data, 'getlist') else []
        vodovod_data = form_data.getlist('vodovod[]') if hasattr(form_data, 'getlist') else []
        gpp_data = form_data.getlist('gpp[]') if hasattr(form_data, 'getlist') else []
        pulpovod_data = form_data.getlist('pulpovod[]') if hasattr(form_data, 'getlist') else []
        raspred_data = form_data.getlist('raspred[]') if hasattr(form_data, 'getlist') else []
        
        if equipment_names and equipment_names[0]:
            # Создаем таблицу
            table = self.doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # Заголовки
            headers = [
                'Техника',
                'Пионерная дамба',
                'Водовод оборотной воды',
                'ГПП-1 ПС 110/10кВ',
                'Магистральный пульповод',
                'Распределительный пульповод'
            ]
            
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Добавляем данные
            for i in range(len(equipment_names)):
                if equipment_names[i]:  # Проверяем что техника не пустая
                    row = table.add_row()
                    row.cells[0].text = equipment_names[i]
                    row.cells[1].text = damba_data[i] if i < len(damba_data) else ''
                    row.cells[2].text = vodovod_data[i] if i < len(vodovod_data) else ''
                    row.cells[3].text = gpp_data[i] if i < len(gpp_data) else ''
                    row.cells[4].text = pulpovod_data[i] if i < len(pulpovod_data) else ''
                    row.cells[5].text = raspred_data[i] if i < len(raspred_data) else ''
    
    def add_works_section(self, form_data):
        """Добавление секции с выполняемыми работами"""
        self.doc.add_heading('Прогресс / Выполняемые работы', level=1)
        
        # Получаем данные о работах
        areas = form_data.getlist('area[]') if hasattr(form_data, 'getlist') else []
        ch_from = form_data.getlist('ch_from[]') if hasattr(form_data, 'getlist') else []
        ch_to = form_data.getlist('ch_to[]') if hasattr(form_data, 'getlist') else []
        work_types = form_data.getlist('work_type[]') if hasattr(form_data, 'getlist') else []
        work_descriptions = form_data.getlist('work_description[]') if hasattr(form_data, 'getlist') else []
        
        if areas and areas[0]:
            table = self.doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # Заголовки
            headers = ['№', 'Участок', 'ПК от', 'ПК до', 'Вид работы', 'Описание работ']
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Добавляем работы
            for i, area in enumerate(areas):
                if area:  # Проверяем что участок не пустой
                    row = table.add_row()
                    row.cells[0].text = str(i + 1)
                    row.cells[1].text = area
                    row.cells[2].text = ch_from[i] if i < len(ch_from) else ''
                    row.cells[3].text = ch_to[i] if i < len(ch_to) else ''
                    row.cells[4].text = work_types[i] if i < len(work_types) else ''
                    row.cells[5].text = work_descriptions[i] if i < len(work_descriptions) else ''
    
    def add_materials_section(self, report_data):
        """Добавление секции поставки материалов"""
        materials_info = report_data.get('materials_delivery', '')
        if materials_info:
            self.doc.add_heading('Поставка материалов', level=1)
            self.doc.add_paragraph(materials_info)
    
    def add_quality_control_section(self, report_data):
        """Добавление секции контроля качества"""
        self.doc.add_heading('Контроль качества', level=1)
        
        quality_sections = [
            ('Западный участок пионерной дамбы:', 'quality_west'),
            ('Северный участок пионерной дамбы:', 'quality_north'),
            ('Магистральный пульповод:', 'quality_pipeline'),
            ('Водовод оборотной воды:', 'quality_water')
        ]
        
        for title, key in quality_sections:
            content = report_data.get(key, '')
            if content:
                p = self.doc.add_paragraph()
                p.add_run(title).bold = True
                self.doc.add_paragraph(content)
    
    def add_remarks_section(self, report_data):
        """Добавление секции замечаний"""
        remarks = report_data.get('remarks', '')
        if remarks:
            self.doc.add_heading('Замечания подрядной организации', level=1)
            self.doc.add_paragraph(remarks)
    
    def add_photos_section_base64(self, form_data):
        """Обработка фотографий из base64 данных"""
        # Ищем все поля с фотографиями
        photo_data_fields = {k: v for k, v in form_data.items() if k.startswith('photo_data_')}
        photo_captions = form_data.getlist('photo_captions[]') if hasattr(form_data, 'getlist') else []
    
        if photo_data_fields:
         self.doc.add_heading('Фото отчеты', level=1)
        
        for i, (field_name, base64_data) in enumerate(photo_data_fields.items()):
            try:
                # Извлекаем base64 данные (убираем префикс data:image/...)
                if ',' in base64_data:
                    base64_content = base64_data.split(',')[1]
                else:
                    base64_content = base64_data
                
                # Декодируем и создаем временный файл
                import base64
                image_data = base64.b64decode(base64_content)
                
                with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as temp_file:
                    temp_file.write(image_data)
                    temp_file_path = temp_file.name
                
                # Добавляем в документ
                self.doc.add_picture(temp_file_path, width=Inches(4))
                
                # Добавляем подпись
                caption = photo_captions[i] if i < len(photo_captions) and photo_captions[i] else f'Фото {i + 1}'
                caption_p = self.doc.add_paragraph(f'Рис. {i + 1}. {caption}')
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_p.runs[0].font.italic = True
                
                self.doc.add_paragraph()
                
                # Удаляем временный файл
                os.unlink(temp_file_path)
                
            except Exception as e:
                print(f"Ошибка при обработке фото {i + 1}: {e}")
                self.doc.add_paragraph(f'[Фото {i + 1} - ошибка загрузки]')
        else:
        # Если нет фотографий, проверяем подписи
         if photo_captions and any(caption for caption in photo_captions):
            self.doc.add_heading('Фото отчеты', level=1)
            self.doc.add_paragraph(f'Подготовлено {len([c for c in photo_captions if c])} фотографий.')
    
    def convert_form_to_word(self, form_data, request_files=None):
        """Основная функция конвертации формы в Word"""
        
        # Создаем словарь из данных формы
        report_data = {}
        for key in form_data.keys():
            if not key.endswith('[]'):  # Не массивы
                report_data[key] = form_data.get(key)
        
        # Добавляем секции в документ
        self.add_header(report_data)
        self.doc.add_page_break()
        
        self.add_equipment_section(form_data)
        self.doc.add_paragraph()
        
        self.add_works_section(form_data)
        self.doc.add_paragraph()
        
        self.add_materials_section(report_data)
        self.add_quality_control_section(report_data)
        # Добавляем фотографии
        self.add_photos_section_base64(form_data)
        self.add_remarks_section(report_data)
        
        # Добавляем фотографии если есть
        if request_files:
            photo_captions = form_data.getlist('photo_captions[]') if hasattr(form_data, 'getlist') else []
            self.add_photos_section(request_files, photo_captions)
        
        # Сохраняем документ
        date_str = report_data.get('report_date', datetime.datetime.now().strftime('%Y-%m-%d'))
        author = report_data.get('author', 'report').replace(' ', '_')
        filename = f"daily_report_{date_str}_{author}.docx"
        
        # Создаем BytesIO для возврата файла
        output = io.BytesIO()
        self.doc.save(output)
        output.seek(0)
        
        return output, filename


def test_word_export():
    """Тестирование экспорта в Word"""
    from werkzeug.datastructures import MultiDict
    
    # Тестовые данные
    test_data = MultiDict([
        ('report_date', '2025-01-15'),
        ('project_name', 'Строительство Хвостового Хозяйства для комплекса МОФ-3'),
        ('location', 'Алмалык'),
        ('author', 'Адилжанов М.А'),
        ('weather', 'Без осадков, 15°C'),
        ('equipment_name[]', 'Каток'),
        ('equipment_name[]', 'Экскаватор'),
        ('damba[]', '2'),
        ('damba[]', '1'),
        ('vodovod[]', '0'),
        ('vodovod[]', '1'),
        ('area[]', 'ВзиС-2'),
        ('ch_from[]', '21+00'),
        ('ch_to[]', '22+00'),
        ('work_type[]', 'Земляные работы'),
        ('work_description[]', 'Описание выполняемых работ'),
        ('materials_delivery', 'Отвалы грунта (супесь): Отвалы подготовки грунта дамбы проверены визуально'),
        ('quality_west', 'Контроль качества западного участка'),
        ('remarks', 'Замечания подрядной организации')
    ])
    
    exporter = DailyReportWordExporter()
    output, filename = exporter.convert_form_to_word(test_data)
    
    # Сохраняем файл
    with open(filename, 'wb') as f:
        f.write(output.getvalue())
    
    print(f"Тестовый Word документ создан: {filename}")
    return True

if __name__ == '__main__':
    test_word_export()